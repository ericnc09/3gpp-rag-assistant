"""
Unified document processor for 3GPP specifications

Automatically detects and processes:
- PDF files (.pdf)
- Modern Word documents (.docx)
- Legacy Word documents (.doc)

This is a convenience wrapper that routes to the appropriate processor
based on file extension.
"""

from pathlib import Path
from typing import List, Dict, Optional, Tuple
import logging
from dataclasses import dataclass
import re
import sys

logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    """Represents a chunk of text from a document"""
    text: str
    metadata: Dict
    chunk_id: int
    
    def to_dict(self) -> Dict:
        """Convert chunk to dictionary format"""
        return {
            "text": self.text,
            "metadata": self.metadata,
            "chunk_id": self.chunk_id
        }


class UnifiedDocumentProcessor:
    """
    Unified processor for PDF, DOCX, and DOC files
    
    Automatically routes to the appropriate processor based on file extension
    """
    
    def __init__(
        self, 
        chunk_size: int = 1000, 
        chunk_overlap: int = 200,
        min_chunk_size: int = 100
    ):
        """
        Initialize unified document processor
        
        Args:
            chunk_size: Target size of each chunk in characters
            chunk_overlap: Number of characters to overlap between chunks
            min_chunk_size: Minimum size for a valid chunk
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size
        
        # Track which processors are available
        self.available_processors = self._detect_processors()
        logger.info(f"Available processors: {list(self.available_processors.keys())}")
    
    def _detect_processors(self) -> Dict[str, bool]:
        """Detect which file format processors are available"""
        processors = {}
        
        # Check for PDF support
        try:
            from pypdf import PdfReader
            processors['pdf'] = True
        except ImportError:
            processors['pdf'] = False
            logger.warning("PDF support not available (install pypdf)")
        
        # Check for DOCX support
        try:
            from docx import Document
            processors['docx'] = True
        except ImportError:
            processors['docx'] = False
            logger.warning("DOCX support not available (install python-docx)")
        
        # Check for DOC support (any method)
        doc_methods = []
        
        # Check antiword
        try:
            import subprocess
            subprocess.run(['antiword', '-v'], capture_output=True, check=True)
            doc_methods.append('antiword')
        except (subprocess.CalledProcessError, FileNotFoundError):
            pass
        
        # Check textract
        try:
            import textract
            doc_methods.append('textract')
        except ImportError:
            pass
        
        # Check win32com
        if sys.platform == 'win32':
            try:
                import win32com.client
                doc_methods.append('win32com')
            except ImportError:
                pass
        
        processors['doc'] = len(doc_methods) > 0
        processors['doc_methods'] = doc_methods
        
        return processors
    
    def _load_pdf(self, file_path: Path) -> Tuple[str, Dict]:
        """Load PDF file"""
        from pypdf import PdfReader
        
        logger.info(f"Loading PDF: {file_path}")
        reader = PdfReader(str(file_path))
        
        metadata = {
            "source": str(file_path.name),
            "num_pages": len(reader.pages),
            "file_path": str(file_path),
            "format": "pdf"
        }
        
        if reader.metadata:
            metadata.update({
                "title": reader.metadata.get("/Title", ""),
                "author": reader.metadata.get("/Author", ""),
            })
        
        text_parts = []
        for page_num, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"\n[Page {page_num}]\n{page_text}")
        
        full_text = "\n".join(text_parts)
        logger.info(f"Extracted {len(full_text)} characters from PDF")
        
        return full_text, metadata
    
    def _load_docx(self, file_path: Path) -> Tuple[str, Dict]:
        """Load DOCX file"""
        from docx import Document
        
        logger.info(f"Loading DOCX: {file_path}")
        doc = Document(str(file_path))
        
        core_props = doc.core_properties
        metadata = {
            "source": str(file_path.name),
            "file_path": str(file_path),
            "format": "docx",
            "title": core_props.title or "",
            "author": core_props.author or "",
            "num_paragraphs": len(doc.paragraphs),
            "num_tables": len(doc.tables)
        }
        
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                if paragraph.style.name.startswith('Heading'):
                    text_parts.append(f"\n## {paragraph.text}\n")
                else:
                    text_parts.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        full_text = "\n".join(text_parts)
        logger.info(f"Extracted {len(full_text)} characters from DOCX")
        
        return full_text, metadata
    
    def _load_doc(self, file_path: Path) -> Tuple[str, Dict]:
        """Load legacy DOC file"""
        logger.info(f"Loading legacy DOC: {file_path}")
        
        # Try antiword first (best for Linux/Mac)
        if 'antiword' in self.available_processors.get('doc_methods', []):
            import subprocess
            result = subprocess.run(
                ['antiword', str(file_path)],
                capture_output=True,
                text=True,
                check=True
            )
            text = result.stdout
            method = 'antiword'
        
        # Try textract
        elif 'textract' in self.available_processors.get('doc_methods', []):
            import textract
            text_bytes = textract.process(str(file_path))
            text = text_bytes.decode('utf-8', errors='ignore')
            method = 'textract'
        
        # Try win32com (Windows)
        elif 'win32com' in self.available_processors.get('doc_methods', []):
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            try:
                doc = word.Documents.Open(str(file_path.resolve()))
                text = doc.Content.Text
                doc.Close(False)
            finally:
                word.Quit()
            method = 'win32com'
        
        else:
            raise RuntimeError(
                "No method available to process .doc files. "
                "Install antiword, textract, or pywin32 (Windows)."
            )
        
        metadata = {
            "source": str(file_path.name),
            "file_path": str(file_path),
            "format": "doc",
            "extraction_method": method
        }
        
        logger.info(f"Extracted {len(text)} characters from DOC using {method}")
        
        return text, metadata
    
    def load_document(self, file_path: Path) -> Tuple[str, Dict]:
        """
        Load any supported document format
        
        Automatically detects format and uses appropriate loader
        
        Args:
            file_path: Path to document file
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = file_path.suffix.lower()
        
        # Route to appropriate loader
        if file_ext == '.pdf':
            if not self.available_processors.get('pdf'):
                raise RuntimeError("PDF support not available. Install: pip install pypdf")
            return self._load_pdf(file_path)
        
        elif file_ext == '.docx':
            if not self.available_processors.get('docx'):
                raise RuntimeError("DOCX support not available. Install: pip install python-docx")
            return self._load_docx(file_path)
        
        elif file_ext == '.doc':
            if not self.available_processors.get('doc'):
                raise RuntimeError(
                    "DOC support not available. Install antiword, textract, or pywin32"
                )
            return self._load_doc(file_path)
        
        else:
            raise ValueError(
                f"Unsupported file format: {file_ext}\n"
                f"Supported formats: .pdf, .docx, .doc"
            )
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        text = re.sub(r'\n\d+\s+3GPP.*?\n', '\n', text)
        text = re.sub(r'[^\w\s.,;:()\[\]\-/\n#|]', '', text)
        return text.strip()
    
    def chunk_text(self, text: str, metadata: Dict) -> List[DocumentChunk]:
        """Split text into overlapping chunks"""
        text = self.clean_text(text)
        
        chunks = []
        chunk_id = 0
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            
            if end < len(text):
                search_start = max(start, end - 100)
                search_text = text[search_start:end + 100]
                sentence_endings = [m.end() for m in re.finditer(r'[.!?]\s+', search_text)]
                
                if sentence_endings:
                    last_sentence = sentence_endings[-1]
                    end = search_start + last_sentence
            
            chunk_text = text[start:end].strip()
            
            if len(chunk_text) >= self.min_chunk_size:
                chunk_metadata = metadata.copy()
                chunk_metadata.update({
                    "chunk_index": chunk_id,
                    "start_char": start,
                    "end_char": end,
                    "chunk_size": len(chunk_text)
                })
                
                chunk = DocumentChunk(
                    text=chunk_text,
                    metadata=chunk_metadata,
                    chunk_id=chunk_id
                )
                chunks.append(chunk)
                chunk_id += 1
            
            start = end - self.chunk_overlap
            if start <= 0:
                start = 1
        
        logger.info(f"Created {len(chunks)} chunks from document")
        return chunks
    
    def process_document(self, file_path: Path) -> List[DocumentChunk]:
        """Process any supported document file"""
        text, metadata = self.load_document(file_path)
        chunks = self.chunk_text(text, metadata)
        return chunks
    
    def process_directory(self, directory: Path) -> List[DocumentChunk]:
        """
        Process all supported documents in a directory
        
        Handles .pdf, .docx, and .doc files
        """
        logger.info(f"Processing directory: {directory}")
        
        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory}")
        
        # Find all supported files
        pdf_files = list(directory.glob("*.pdf"))
        docx_files = list(directory.glob("*.docx"))
        doc_files = [f for f in directory.glob("*.doc") if f.suffix.lower() == '.doc']
        
        all_files = pdf_files + docx_files + doc_files
        
        if not all_files:
            logger.warning(f"No supported files found in {directory}")
            logger.info("Looking for: .pdf, .docx, .doc files")
            return []
        
        logger.info(f"Found {len(all_files)} files: {len(pdf_files)} PDF, {len(docx_files)} DOCX, {len(doc_files)} DOC")
        
        all_chunks = []
        failed_files = []
        
        for doc_file in all_files:
            try:
                chunks = self.process_document(doc_file)
                all_chunks.extend(chunks)
                logger.info(f"✅ {doc_file.name}: {len(chunks)} chunks")
            except Exception as e:
                logger.error(f"❌ {doc_file.name}: {e}")
                failed_files.append(doc_file.name)
                continue
        
        logger.info(f"Total chunks: {len(all_chunks)}")
        
        if failed_files:
            logger.warning(f"Failed: {', '.join(failed_files)}")
        
        return all_chunks
    
    def save_chunks(self, chunks: List[DocumentChunk], output_path: Path) -> None:
        """Save chunks to JSON file"""
        import json
        
        chunks_dict = [chunk.to_dict() for chunk in chunks]
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(chunks_dict, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Saved {len(chunks)} chunks to {output_path}")


def main():
    """Example usage"""
    import sys
    from pathlib import Path
    
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s'
    )
    
    print(f"\n{'='*60}")
    print("3GPP Unified Document Processor")
    print("Supports: PDF (.pdf), Word (.docx), Legacy Word (.doc)")
    print(f"{'='*60}\n")
    
    processor = UnifiedDocumentProcessor(
        chunk_size=1000,
        chunk_overlap=200
    )
    
    # Show capabilities
    print("Available format support:")
    print(f"  PDF:  {'✅' if processor.available_processors.get('pdf') else '❌'}")
    print(f"  DOCX: {'✅' if processor.available_processors.get('docx') else '❌'}")
    print(f"  DOC:  {'✅' if processor.available_processors.get('doc') else '❌'}")
    if processor.available_processors.get('doc'):
        methods = processor.available_processors.get('doc_methods', [])
        print(f"        Methods: {', '.join(methods)}")
    print()
    
    data_dir = Path("data/raw")
    
    if not data_dir.exists():
        print(f"Error: {data_dir} does not exist")
        sys.exit(1)
    
    try:
        chunks = processor.process_directory(data_dir)
        
        if not chunks:
            print("No documents processed.")
            sys.exit(1)
        
        output_file = Path("data/processed/chunks.json")
        processor.save_chunks(chunks, output_file)
        
        print(f"\n{'='*60}")
        print("Summary:")
        print(f"{'='*60}")
        print(f"Total chunks: {len(chunks)}")
        print(f"Output: {output_file}")
        
        # Statistics by format
        formats = {}
        for chunk in chunks:
            fmt = chunk.metadata.get('format', 'unknown')
            formats[fmt] = formats.get(fmt, 0) + 1
        
        print(f"\nBy format:")
        for fmt, count in sorted(formats.items()):
            print(f"  {fmt.upper()}: {count} chunks")
        
        print(f"\n{'='*60}")
        print("✅ Complete!")
        print(f"{'='*60}\n")
        
    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
