"""
Document processing module for parsing and chunking 3GPP Word documents (.doc/.docx)

This module handles:
- DOCX/DOC text extraction
- Text cleaning and preprocessing
- Intelligent chunking with overlap
- Metadata extraction
"""

from pathlib import Path
from typing import List, Dict, Optional
import logging
from dataclasses import dataclass
import re

try:
    from docx import Document
except ImportError:
    print("Warning: python-docx not installed. Run: pip install python-docx")
    Document = None

logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    """Represents a chunk of text from a document"""
    text: str
    metadata: Dict
    chunk_id: int
    
    def to_dict(self) -> Dict:
        """Convert chunk to dictionary format"""
        return {
            "text": self.text,
            "metadata": self.metadata,
            "chunk_id": self.chunk_id
        }


class DocumentProcessor:
    """Process 3GPP Word documents (.doc/.docx) into chunks for RAG"""
    
    def __init__(
        self, 
        chunk_size: int = 1000, 
        chunk_overlap: int = 200,
        min_chunk_size: int = 100
    ):
        """
        Initialize document processor
        
        Args:
            chunk_size: Target size of each chunk in characters
            chunk_overlap: Number of characters to overlap between chunks
            min_chunk_size: Minimum size for a valid chunk
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size
        
        if not Document:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
    
    def load_docx(self, file_path: Path) -> tuple[str, Dict]:
        """
        Load and extract text from a DOCX file
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        logger.info(f"Loading DOCX: {file_path}")
        
        if not file_path.exists():
            raise FileNotFoundError(f"Document file not found: {file_path}")
        
        try:
            doc = Document(str(file_path))
            
            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                "source": str(file_path.name),
                "file_path": str(file_path),
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else ""
            }
            
            # Extract text from paragraphs
            text_parts = []
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    # Keep track of headings vs body text
                    if paragraph.style.name.startswith('Heading'):
                        text_parts.append(f"\n## {paragraph.text}\n")
                    else:
                        text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            full_text = "\n".join(text_parts)
            
            # Add paragraph count to metadata
            metadata["num_paragraphs"] = len(doc.paragraphs)
            metadata["num_tables"] = len(doc.tables)
            
            logger.info(f"Extracted {len(full_text)} characters from {len(doc.paragraphs)} paragraphs")
            
            return full_text, metadata
            
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
            raise
    
    def load_doc(self, file_path: Path) -> tuple[str, Dict]:
        """
        Load and extract text from a legacy .doc file
        
        Note: .doc files (MS Word 97-2003) require additional libraries.
        This method attempts to convert .doc to .docx first.
        
        Args:
            file_path: Path to the .doc file
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        logger.warning(f"Legacy .doc format detected: {file_path}")
        logger.warning("For best results, convert .doc files to .docx format manually")
        
        # Try using antiword if available (Linux/Mac)
        try:
            import subprocess
            result = subprocess.run(
                ['antiword', str(file_path)],
                capture_output=True,
                text=True,
                check=True
            )
            text = result.stdout
            
            metadata = {
                "source": str(file_path.name),
                "file_path": str(file_path),
                "format": "doc (legacy)",
                "conversion_method": "antiword"
            }
            
            logger.info(f"Extracted {len(text)} characters using antiword")
            return text, metadata
            
        except (subprocess.CalledProcessError, FileNotFoundError):
            logger.error("antiword not available. Please install antiword or convert .doc to .docx")
            logger.info("On Ubuntu: sudo apt-get install antiword")
            logger.info("On Mac: brew install antiword")
            logger.info("Or convert manually: Open in Word → Save As → .docx")
            raise RuntimeError("Cannot process .doc files. Please convert to .docx format.")
    
    def load_document(self, file_path: Path) -> tuple[str, Dict]:
        """
        Load document regardless of .doc or .docx format
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        file_ext = file_path.suffix.lower()
        
        if file_ext == '.docx':
            return self.load_docx(file_path)
        elif file_ext == '.doc':
            return self.load_doc(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}. Use .doc or .docx")
    
    def clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text
        
        Args:
            text: Raw text from document
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Normalize line breaks
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Remove page numbers and headers/footers (common patterns in 3GPP docs)
        text = re.sub(r'\n\d+\s+3GPP.*?\n', '\n', text)
        text = re.sub(r'^\s*\d+\s*$', '', text, flags=re.MULTILINE)
        
        # Remove special Unicode characters but keep basic punctuation
        text = re.sub(r'[^\w\s.,;:()\[\]\-/\n#|]', '', text)
        
        # Clean up table separators
        text = re.sub(r'\s*\|\s*\|\s*', ' | ', text)
        
        return text.strip()
    
    def chunk_text(self, text: str, metadata: Dict) -> List[DocumentChunk]:
        """
        Split text into overlapping chunks
        
        Args:
            text: Full document text
            metadata: Document metadata
            
        Returns:
            List of DocumentChunk objects
        """
        # Clean the text first
        text = self.clean_text(text)
        
        chunks = []
        chunk_id = 0
        start = 0
        
        while start < len(text):
            # Calculate end position
            end = start + self.chunk_size
            
            # If this is not the last chunk, try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings near the chunk boundary
                search_start = max(start, end - 100)
                search_text = text[search_start:end + 100]
                
                # Find the last sentence ending
                sentence_endings = [m.end() for m in re.finditer(r'[.!?]\s+', search_text)]
                
                if sentence_endings:
                    # Adjust end to the sentence boundary
                    last_sentence = sentence_endings[-1]
                    end = search_start + last_sentence
            
            # Extract chunk text
            chunk_text = text[start:end].strip()
            
            # Only create chunk if it meets minimum size
            if len(chunk_text) >= self.min_chunk_size:
                # Create chunk metadata
                chunk_metadata = metadata.copy()
                chunk_metadata.update({
                    "chunk_index": chunk_id,
                    "start_char": start,
                    "end_char": end,
                    "chunk_size": len(chunk_text)
                })
                
                chunk = DocumentChunk(
                    text=chunk_text,
                    metadata=chunk_metadata,
                    chunk_id=chunk_id
                )
                chunks.append(chunk)
                chunk_id += 1
            
            # Move start position with overlap
            start = end - self.chunk_overlap
            
            # Prevent infinite loop
            if start <= 0:
                start = 1
        
        logger.info(f"Created {len(chunks)} chunks from document")
        
        return chunks
    
    def process_document(self, file_path: Path) -> List[DocumentChunk]:
        """
        Complete pipeline: Load document → Extract text → Create chunks
        
        Args:
            file_path: Path to document file (.doc or .docx)
            
        Returns:
            List of document chunks ready for embedding
        """
        # Load document
        text, metadata = self.load_document(file_path)
        
        # Create chunks
        chunks = self.chunk_text(text, metadata)
        
        return chunks
    
    def process_directory(self, directory: Path) -> List[DocumentChunk]:
        """
        Process all Word documents in a directory
        
        Args:
            directory: Path to directory containing .doc/.docx files
            
        Returns:
            List of all chunks from all documents
        """
        logger.info(f"Processing directory: {directory}")
        
        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory}")
        
        # Find all Word document files
        doc_files = list(directory.glob("*.doc")) + list(directory.glob("*.docx"))
        
        if not doc_files:
            logger.warning(f"No .doc or .docx files found in {directory}")
            return []
        
        logger.info(f"Found {len(doc_files)} Word document files")
        
        all_chunks = []
        for doc_file in doc_files:
            try:
                chunks = self.process_document(doc_file)
                all_chunks.extend(chunks)
                logger.info(f"Processed {doc_file.name}: {len(chunks)} chunks")
            except Exception as e:
                logger.error(f"Error processing {doc_file.name}: {e}")
                continue
        
        logger.info(f"Total chunks created: {len(all_chunks)}")
        
        return all_chunks
    
    def save_chunks(self, chunks: List[DocumentChunk], output_path: Path) -> None:
        """
        Save chunks to JSON file
        
        Args:
            chunks: List of document chunks
            output_path: Path to save JSON file
        """
        import json
        
        chunks_dict = [chunk.to_dict() for chunk in chunks]
        
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(chunks_dict, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Saved {len(chunks)} chunks to {output_path}")


def main():
    """Example usage of DocumentProcessor"""
    import sys
    from pathlib import Path
    
    # Set up logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    # Initialize processor
    processor = DocumentProcessor(
        chunk_size=1000,
        chunk_overlap=200
    )
    
    # Process documents
    data_dir = Path("data/raw")
    
    if not data_dir.exists():
        print(f"Error: Directory {data_dir} does not exist")
        print("Please create it and add your 3GPP Word documents (.doc or .docx)")
        sys.exit(1)
    
    print(f"\n{'='*60}")
    print("3GPP Word Document Processor")
    print(f"{'='*60}\n")
    
    try:
        # Process all Word documents
        chunks = processor.process_directory(data_dir)
        
        if not chunks:
            print("\nNo documents were processed. Please add .doc or .docx files to data/raw/")
            sys.exit(1)
        
        # Save processed chunks
        output_file = Path("data/processed/chunks.json")
        processor.save_chunks(chunks, output_file)
        
        # Print summary
        print(f"\n{'='*60}")
        print("Processing Summary:")
        print(f"{'='*60}")
        print(f"Total chunks created: {len(chunks)}")
        print(f"Output saved to: {output_file}")
        
        # Show sample chunk
        if chunks:
            print(f"\nSample chunk:")
            print(f"  Source: {chunks[0].metadata['source']}")
            print(f"  Chunk size: {len(chunks[0].text)} characters")
            print(f"  Preview: {chunks[0].text[:200]}...")
        
        # Show statistics
        sources = set(chunk.metadata['source'] for chunk in chunks)
        print(f"\nProcessed documents:")
        for source in sorted(sources):
            source_chunks = [c for c in chunks if c.metadata['source'] == source]
            print(f"  - {source}: {len(source_chunks)} chunks")
        
        print(f"\n{'='*60}")
        print("✅ Processing complete!")
        print(f"{'='*60}\n")
        
    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
